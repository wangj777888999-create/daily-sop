import re
import os
import hashlib
from typing import List
from .models import ParsedDocument


def compute_content_hash(file_path: str) -> str:
    h = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def parse_document(file_path: str, doc_type: str) -> ParsedDocument:
    ext = os.path.splitext(file_path)[1].lower()
    doc_type = doc_type.upper()

    if doc_type == "PDF" or ext == ".pdf":
        return _parse_pdf(file_path)
    elif doc_type == "DOCX" or ext == ".docx":
        return _parse_docx(file_path)
    elif doc_type == "XLSX" or ext in (".xlsx", ".xls"):
        return _parse_xlsx(file_path)
    elif doc_type == "MD" or ext == ".md":
        return _parse_md(file_path)
    elif doc_type == "TXT" or ext == ".txt":
        return _parse_txt(file_path)
    else:
        raise ValueError(f"Unsupported document type: {doc_type}")


def _parse_pdf(path: str) -> ParsedDocument:
    import pdfplumber

    chunks = []
    full_lines = []

    with pdfplumber.open(path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            text = page.extract_text()
            if not text:
                continue
            paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
            for para in paragraphs:
                chunk_type = "heading" if _looks_like_heading(para) else "paragraph"
                full_lines.append(para)
                chunks.append({
                    "chunk_type": chunk_type,
                    "level": 1 if chunk_type == "heading" else 0,
                    "text": para,
                    "page": page_num,
                    "heading_path": para if chunk_type == "heading" else "",
                })

    return ParsedDocument(
        doc_id="",
        full_text="\n".join(full_lines),
        chunks=chunks,
    )


def _parse_docx(path: str) -> ParsedDocument:
    from docx import Document

    doc = Document(path)
    chunks = []
    full_lines = []
    heading_stack = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style = para.style.name if para.style else ""
        chunk_type = "paragraph"
        level = 0

        if style.startswith("Heading") or style.startswith("heading"):
            chunk_type = "heading"
            try:
                level = int(style.split()[-1])
            except ValueError:
                level = 1
            while len(heading_stack) >= level:
                heading_stack.pop()
            heading_stack.append(text)
        elif style.startswith("List"):
            chunk_type = "list_item"

        heading_path = " > ".join(heading_stack) if heading_stack else ""

        full_lines.append(text)
        chunks.append({
            "chunk_type": chunk_type,
            "level": level,
            "text": text,
            "page": 0,
            "heading_path": heading_path,
        })

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        table_text = "\n".join(rows)
        if table_text.strip():
            full_lines.append(table_text)
            chunks.append({
                "chunk_type": "table",
                "level": 0,
                "text": table_text,
                "page": 0,
                "heading_path": heading_stack[-1] if heading_stack else "",
            })

    return ParsedDocument(
        doc_id="",
        full_text="\n".join(full_lines),
        chunks=chunks,
    )


def _parse_xlsx(path: str) -> ParsedDocument:
    import openpyxl

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    chunks = []
    full_lines = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(max_row=500, values_only=True))
        if not rows:
            continue

        header = [str(c) if c is not None else "" for c in rows[0]]
        header_text = " | ".join(header)
        full_lines.append(f"[Sheet: {sheet_name}]")
        full_lines.append(header_text)

        chunks.append({
            "chunk_type": "table",
            "level": 0,
            "text": f"Sheet: {sheet_name}\nColumns: {header_text}",
            "page": 0,
            "heading_path": sheet_name,
        })

        sample_text = f"Sheet: {sheet_name}\n{header_text}"
        for row in rows[1:51]:
            row_text = " | ".join([str(c) if c is not None else "" for c in row])
            full_lines.append(row_text)
            sample_text += "\n" + row_text

        chunks.append({
            "chunk_type": "paragraph",
            "level": 0,
            "text": sample_text[:1000],
            "page": 0,
            "heading_path": sheet_name,
        })

    wb.close()
    return ParsedDocument(
        doc_id="",
        full_text="\n".join(full_lines),
        chunks=chunks,
    )


def _parse_txt(path: str) -> ParsedDocument:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()

    lines = text.split("\n")
    chunks = []
    full_lines = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        chunk_type = "heading" if _looks_like_heading(stripped) else "paragraph"
        full_lines.append(stripped)
        chunks.append({
            "chunk_type": chunk_type,
            "level": 1 if chunk_type == "heading" else 0,
            "text": stripped,
            "page": 0,
            "heading_path": stripped if chunk_type == "heading" else "",
        })

    return ParsedDocument(
        doc_id="",
        full_text="\n".join(full_lines),
        chunks=chunks,
    )


def _parse_md(path: str) -> ParsedDocument:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()

    lines = text.split("\n")
    chunks = []
    full_lines = []
    heading_stack = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.+)", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2)
            while len(heading_stack) >= level:
                heading_stack.pop()
            heading_stack.append(heading_text)
            full_lines.append(stripped)
            chunks.append({
                "chunk_type": "heading",
                "level": level,
                "text": heading_text,
                "page": 0,
                "heading_path": " > ".join(heading_stack),
            })
        elif stripped.startswith("```"):
            continue
        elif re.match(r"^[-*+]\s", stripped):
            full_lines.append(stripped)
            chunks.append({
                "chunk_type": "list_item",
                "level": 0,
                "text": stripped,
                "page": 0,
                "heading_path": " > ".join(heading_stack) if heading_stack else "",
            })
        else:
            full_lines.append(stripped)
            chunks.append({
                "chunk_type": "paragraph",
                "level": 0,
                "text": stripped,
                "page": 0,
                "heading_path": " > ".join(heading_stack) if heading_stack else "",
            })

    return ParsedDocument(
        doc_id="",
        full_text="\n".join(full_lines),
        chunks=chunks,
    )


def _looks_like_heading(text: str) -> bool:
    """Heuristic: short line ending without punctuation, or all-caps, or numbered heading pattern."""
    if len(text) > 60:
        return False
    if re.match(r"^第[一二三四五六七八九十百千万\d]+\s*[章节条款]", text):
        return True
    if re.match(r"^\d+[\.、\s)]", text):
        return True
    if text.isupper() and len(text) < 40:
        return True
    if not re.search(r"[。，；：！？、,;:!?.]$", text) and len(text) < 30:
        return True
    return False
